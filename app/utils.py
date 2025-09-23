import logging
import re
from docxtpl import DocxTemplate, Subdoc
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any

logger = logging.getLogger(__name__)

def _add_runs_with_bolding(paragraph, text_content: str):
    """
    Parses a string containing <b>...</b> tags and adds formatted runs
    to the given paragraph object.
    """
    parts = re.split(r'(<b>|</b>)', text_content)
    
    is_bold = False
    for part in parts:
        if part == '<b>':
            is_bold = True
            continue
        elif part == '</b>':
            is_bold = False
            continue
        
        if part: 
            run = paragraph.add_run(part)
            run.bold = is_bold

def create_subdoc_from_structured_data(template: DocxTemplate, structured_data: List[Dict[str, Any]]) -> Subdoc:
    """
    Takes a list of structured content objects (e.g., from an LLM) and
    builds a docxtpl Subdoc object with appropriate formatting.
    """
    subdoc = template.new_subdoc()

    if not isinstance(structured_data, list):
        logger.warning(f"Expected a list for structured data, but got {type(structured_data)}. Rendering as plain text.")
        p = subdoc.add_paragraph(str(structured_data))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return subdoc

    for item in structured_data:
        content_type = item.get("type", "paragraph")
        content = item.get("content", "")
        
        p = None
        if content_type == "heading":
            p = subdoc.add_heading(level=3)
        
        elif content_type == "bullet":
            p = subdoc.add_paragraph(style='Normal')
            p.add_run('-\t')
        else: 
            p = subdoc.add_paragraph()

        _add_runs_with_bolding(p, content)
        
        if content_type != "heading":
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return subdoc
