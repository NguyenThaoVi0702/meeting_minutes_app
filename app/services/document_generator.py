import json
import logging
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate

from app.utils import create_subdoc_from_structured_data

logger = logging.getLogger(__name__)
TEMPLATES_DIR = Path(__file__).resolve().parent.parent.parent / "templates"

# ===================================================================
#   DOCX Generation for Template-Based Summaries (bbh_hdqt, nghi_quyet)
# ===================================================================
def generate_templated_document(template_type: str, llm_json_output: str) -> BytesIO:
    """
    Generates a DOCX file from a template and JSON data.
    """
    template_map = {
        "bbh_hdqt": "bbh_hdqt_template.docx",
        "nghi_quyet": "nghi_quyet_template.docx",
    }
    template_filename = template_map.get(template_type)
    if not template_filename:
        raise ValueError(f"Invalid template type: {template_type}")

    template_path = TEMPLATES_DIR / template_filename
    if not template_path.exists():
        raise FileNotFoundError(f"Template file not found at: {template_path}")

    try:
        doc = DocxTemplate(template_path)
        context_data = json.loads(llm_json_output)
        
        render_context = {}
        for key, value in context_data.items():
            if isinstance(value, list):
                render_context[key] = create_subdoc_from_structured_data(doc, value)
            else:
                render_context[key] = value
        
        doc.render(render_context)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except json.JSONDecodeError:
        logger.error(f"Failed to decode LLM JSON: {llm_json_output}")
        raise ValueError("AI service returned invalid JSON.")
    except Exception as e:
        logger.error(f"Failed to generate document for template '{template_type}': {e}", exc_info=True)
        raise RuntimeError("An unexpected error occurred during document generation.")


# ===================================================================
#   DOCX Generation for Markdown-Based Summaries (topic, speaker)
# ===================================================================

def _parse_inline_markdown(text: str, paragraph):
    """
    Parses a string for **bold** markdown and adds formatted
    runs to a python-docx paragraph object.
    """
    # Use regex to find all parts that are bold (**) or just text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # It's a bold part, remove the asterisks and add as a bold run
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            # It's a normal text part
            paragraph.add_run(part)

def generate_docx_from_markdown(markdown_content: str) -> BytesIO:
    """
    Converts a markdown string to a well-formatted DOCX file in memory
    using the python-docx library. 
    """
    logger.info("Generating DOCX from markdown using custom python-docx parser...")
    try:
        doc = Document()
        lines = markdown_content.strip().split('\n')

        for line in lines:
            line = line.rstrip()
            if not line.strip():
                # Add an empty paragraph for newlines
                doc.add_paragraph()
                continue

            # Headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                doc.add_heading(text, level=min(level, 4))
            # Unordered list
            elif line.strip().startswith(('-', '*')):
                text = line.strip().lstrip('-* ').strip()
                # Use 'List Bullet' style for proper indentation
                p = doc.add_paragraph(style='List Bullet')
                _parse_inline_markdown(text, p)
            # Ordered list (simple check)
            elif line.strip() and line.strip()[0].isdigit() and '.' in line:
                text = line.split('.', 1)[1].strip()
                p = doc.add_paragraph(style='List Number')
                _parse_inline_markdown(text, p)
            # Normal paragraph
            else:
                p = doc.add_paragraph()
                _parse_inline_markdown(line, p)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info("Custom parser successfully generated DOCX buffer.")
        return buffer
    except Exception as e:
        logger.error(f"Custom DOCX parser failed: {e}", exc_info=True)
        raise RuntimeError(f"Failed to generate DOCX from markdown: {e}")
