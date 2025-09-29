import pypandoc
import os 
from io import BytesIO
import logging
import docx

# Define log
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


from docx import Document
from docx.shared import Pt

def parse_inline_markdown(text, paragraph):
    i = 0
    while i < len(text):
        if text[i:i+2] == '**':  # bold start/end
            i += 2
            end = text.find('**', i)
            if end == -1:
                end = len(text)
            run = paragraph.add_run(text[i:end])
            run.bold = True
            i = end + 2
        elif text[i] == '*':  # italic start/end
            i += 1
            end = text.find('*', i)
            if end == -1:
                end = len(text)
            run = paragraph.add_run(text[i:end])
            run.italic = True
            i = end + 1
        else:
            # add normal text until next * or **
            next_bold = text.find('**', i)
            next_italic = text.find('*', i)
            next_special = min([pos for pos in [next_bold, next_italic] if pos != -1], default=len(text))
            run = paragraph.add_run(text[i:next_special])
            i = next_special
    return paragraph

def md_to_docx(md_path):
    docx_path = 'output.docx'
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_list = False
    list_type = None  # 'ul' or 'ol'

    for line in lines:
        line = line.rstrip('\n').rstrip('\r')

        if line.strip() == '':
            # empty line ends list if any
            in_list = False
            list_type = None
            continue

        # Headings
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            doc.add_heading(text, level=level)
            in_list = False
            list_type = None
        # Unordered list (simple)
        elif line.lstrip().startswith(('-', '*')):
            text = line.lstrip('-* ').strip()
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_markdown(text, p)
            in_list = True
            list_type = 'ul'
        # Ordered list (simple)
        elif line.lstrip()[0:2].isdigit() and line.lstrip()[2:4] == '. ':
            # naive detection for ordered list: starts with digits + dot + space
            text = line.lstrip()[3:].strip()
            p = doc.add_paragraph(style='List Number')
            parse_inline_markdown(text, p)
            in_list = True
            list_type = 'ol'
        else:
            # Normal paragraph
            p = doc.add_paragraph()
            parse_inline_markdown(line.strip(), p)
            in_list = False
            list_type = None

    # doc.save(docx_path)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def process_filename(original_filename):
    """
    Converts a filename like 'Họp UB QLVG_291.wav' to 'Họp UB QLVG.wav'
    and extracts the number '291'.

    Args:
        original_filename (str): The original filename.

    Returns:
        tuple: A tuple containing the new filename (str) and the extracted number (str).
               Returns (None, None) if the format doesn't match.
    """
    base, ext = os.path.splitext(original_filename)
    parts = base.split('_')

    if len(parts) > 1 and parts[-1].isdigit():
        number = parts[-1]
        new_base = '_'.join(parts[:-1])
        new_filename = new_base + ext
        return new_filename, number
    else:
        return original_filename, 0

def save_markdown_to_file(filename: str, content: str):
    """
    Saves markdown content to a specified file.

    Args:
        filename (str): The name of the file to save (e.g., 'my_document.md').
        content (str): The markdown string to write to the file.
    """
    try:
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"Markdown content successfully saved to '{filename}'")
    except IOError as e:
        print(f"Error saving file '{filename}': {e}")

def create_meeting_minutes_doc_buffer(markdown_content: str) -> BytesIO:
    """
    Converts markdown content to a DOCX buffer using pypandoc.
    Returns a BytesIO object containing the DOCX file data.
    """
    markdown_file_path = 'temp.md'
    docx_output_path = 'temp_output.docx'

    save_markdown_to_file(markdown_file_path, markdown_content)

    try:
        logger.info(f"Converting markdown to DOCX using Pandoc...")
        pypandoc.convert_file(
            markdown_file_path,
            to='docx',
            outputfile=docx_output_path,
            extra_args=['--standalone'] # Ensures a complete document
        )
        logger.info("Pandoc conversion successful!")

        # Read the generated DOCX file into a BytesIO buffer
        with open(docx_output_path, "rb") as f:
            buffer = BytesIO(f.read())
        buffer.seek(0) # Rewind the buffer to the beginning for reading

        return buffer

    except Exception as e:
        logger.error(f"Pandoc - An unexpected error occurred during DOCX creation: {e}")
        return md_to_docx(markdown_file_path)
    finally:
        # Clean up temporary files
        if os.path.exists(markdown_file_path):
            os.remove(markdown_file_path)
        if os.path.exists(docx_output_path):
            os.remove(docx_output_path)
        # return None
